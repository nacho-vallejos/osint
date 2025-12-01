"""
Metadata Extraction API Routes
Upload files (images, documents) and extract human-readable metadata
"""

from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from typing import Dict, Any, Optional
import io
import logging

from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import piexif

logger = logging.getLogger(__name__)

router = APIRouter()


def extract_gps_coordinates(gps_info: Dict) -> Optional[Dict[str, float]]:
    """
    Convert GPS EXIF data to decimal coordinates.
    
    Args:
        gps_info: GPS information from EXIF
        
    Returns:
        Dictionary with latitude, longitude, altitude
    """
    try:
        def convert_to_degrees(value):
            """Convert GPS coordinates to degrees"""
            d, m, s = value
            return d + (m / 60.0) + (s / 3600.0)
        
        lat = gps_info.get('GPSLatitude')
        lat_ref = gps_info.get('GPSLatitudeRef')
        lon = gps_info.get('GPSLongitude')
        lon_ref = gps_info.get('GPSLongitudeRef')
        
        if lat and lon and lat_ref and lon_ref:
            latitude = convert_to_degrees(lat)
            if lat_ref == 'S':
                latitude = -latitude
                
            longitude = convert_to_degrees(lon)
            if lon_ref == 'W':
                longitude = -longitude
            
            altitude = gps_info.get('GPSAltitude')
            
            result = {
                'latitude': latitude,
                'longitude': longitude,
                'google_maps_url': f"https://www.google.com/maps?q={latitude},{longitude}",
                'coordinates_string': f"{latitude:.6f}, {longitude:.6f}"
            }
            
            if altitude:
                result['altitude'] = float(altitude)
            
            return result
    except Exception as e:
        logger.error(f"Error converting GPS coordinates: {e}")
        return None


def extract_image_metadata(image_data: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract EXIF metadata from images (JPEG, PNG, etc.)
    
    Args:
        image_data: Binary image data
        filename: Original filename
        
    Returns:
        Dictionary with human-readable metadata
    """
    try:
        image = Image.open(io.BytesIO(image_data))
        
        metadata = {
            'file_type': 'image',
            'filename': filename,
            'format': image.format,
            'size': {
                'width': image.width,
                'height': image.height,
                'megapixels': round((image.width * image.height) / 1_000_000, 2)
            },
            'mode': image.mode,
            'has_exif': False,
            'exif_data': {},
            'gps_location': None
        }
        
        # Extract EXIF data
        exif_data = image._getexif()
        
        if exif_data:
            metadata['has_exif'] = True
            
            # Translate EXIF tags to human-readable names
            human_readable_exif = {}
            gps_info = {}
            
            for tag_id, value in exif_data.items():
                tag_name = TAGS.get(tag_id, tag_id)
                
                # Handle GPS Info separately
                if tag_name == 'GPSInfo':
                    for gps_tag in value:
                        gps_tag_name = GPSTAGS.get(gps_tag, gps_tag)
                        gps_info[gps_tag_name] = value[gps_tag]
                else:
                    # Convert bytes to string for display
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8', errors='ignore')
                        except:
                            value = str(value)
                    # Convert IFDRational to float for JSON serialization
                    elif hasattr(value, 'numerator') and hasattr(value, 'denominator'):
                        value = float(value.numerator) / float(value.denominator) if value.denominator != 0 else 0
                    # Convert tuples of IFDRationals
                    elif isinstance(value, tuple):
                        try:
                            value = tuple(
                                float(v.numerator) / float(v.denominator) if hasattr(v, 'numerator') else v
                                for v in value
                            )
                        except:
                            value = str(value)
                    
                    human_readable_exif[tag_name] = value
            
            metadata['exif_data'] = human_readable_exif
            
            # Extract common fields
            metadata['camera'] = {
                'make': human_readable_exif.get('Make'),
                'model': human_readable_exif.get('Model'),
                'software': human_readable_exif.get('Software')
            }
            
            metadata['capture_info'] = {
                'datetime': human_readable_exif.get('DateTime'),
                'datetime_original': human_readable_exif.get('DateTimeOriginal'),
                'iso': human_readable_exif.get('ISOSpeedRatings'),
                'focal_length': human_readable_exif.get('FocalLength'),
                'exposure_time': human_readable_exif.get('ExposureTime'),
                'f_number': human_readable_exif.get('FNumber'),
                'flash': human_readable_exif.get('Flash')
            }
            
            # Process GPS data if available
            if gps_info:
                gps_coords = extract_gps_coordinates(gps_info)
                if gps_coords:
                    metadata['gps_location'] = gps_coords
                    metadata['has_gps'] = True
                else:
                    metadata['has_gps'] = False
            else:
                metadata['has_gps'] = False
        
        return metadata
        
    except Exception as e:
        logger.error(f"Error extracting image metadata: {e}")
        return {
            'file_type': 'image',
            'filename': filename,
            'error': str(e),
            'success': False
        }


def extract_pdf_metadata(pdf_data: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract metadata from PDF files
    
    Args:
        pdf_data: Binary PDF data
        filename: Original filename
        
    Returns:
        Dictionary with human-readable metadata
    """
    try:
        pdf_reader = PdfReader(io.BytesIO(pdf_data))
        
        metadata = {
            'file_type': 'pdf',
            'filename': filename,
            'num_pages': len(pdf_reader.pages),
            'encrypted': pdf_reader.is_encrypted,
            'metadata': {}
        }
        
        # Extract PDF metadata
        if pdf_reader.metadata:
            info = pdf_reader.metadata
            metadata['metadata'] = {
                'author': info.get('/Author'),
                'creator': info.get('/Creator'),
                'producer': info.get('/Producer'),
                'subject': info.get('/Subject'),
                'title': info.get('/Title'),
                'creation_date': info.get('/CreationDate'),
                'modification_date': info.get('/ModDate')
            }
        
        return metadata
        
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {e}")
        return {
            'file_type': 'pdf',
            'filename': filename,
            'error': str(e),
            'success': False
        }


def extract_docx_metadata(docx_data: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract metadata from DOCX files
    
    Args:
        docx_data: Binary DOCX data
        filename: Original filename
        
    Returns:
        Dictionary with human-readable metadata
    """
    try:
        doc = DocxDocument(io.BytesIO(docx_data))
        core_props = doc.core_properties
        
        metadata = {
            'file_type': 'docx',
            'filename': filename,
            'metadata': {
                'author': core_props.author,
                'last_modified_by': core_props.last_modified_by,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'title': core_props.title,
                'subject': core_props.subject,
                'category': core_props.category,
                'comments': core_props.comments,
                'keywords': core_props.keywords,
                'revision': core_props.revision
            },
            'statistics': {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
        }
        
        return metadata
        
    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {e}")
        return {
            'file_type': 'docx',
            'filename': filename,
            'error': str(e),
            'success': False
        }


@router.post("/extract", response_class=JSONResponse)
async def extract_metadata(file: UploadFile = File(...)):
    """
    Extract metadata from uploaded file (images, PDF, DOCX).
    
    Supports:
    - Images: JPEG, PNG, TIFF (EXIF data including GPS coordinates)
    - PDF: Author, creator, dates
    - DOCX: Author, dates, revision history
    
    Returns human-readable metadata with potential Google Maps link for GPS data.
    """
    try:
        # Read file data
        file_data = await file.read()
        file_size = len(file_data)
        filename = file.filename or "unknown"
        
        # Detect file type
        file_extension = filename.split('.')[-1].lower() if '.' in filename else ''
        
        result = {
            'success': True,
            'filename': filename,
            'file_size_bytes': file_size,
            'file_size_mb': round(file_size / (1024 * 1024), 2),
            'file_extension': file_extension
        }
        
        # Route to appropriate extractor
        if file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'tif']:
            metadata = extract_image_metadata(file_data, filename)
            result['metadata'] = metadata
            
            # Add triangulation suggestion if GPS found
            if metadata.get('has_gps') and metadata.get('gps_location'):
                result['triangulation'] = {
                    'available': True,
                    'type': 'gps_coordinates',
                    'location': metadata['gps_location'],
                    'suggestion': 'GPS coordinates found! Click the Google Maps link to visualize location.'
                }
        
        elif file_extension == 'pdf':
            metadata = extract_pdf_metadata(file_data, filename)
            result['metadata'] = metadata
        
        elif file_extension in ['docx', 'doc']:
            metadata = extract_docx_metadata(file_data, filename)
            result['metadata'] = metadata
        
        else:
            result['success'] = False
            result['error'] = f"Unsupported file type: {file_extension}"
            result['supported_types'] = ['jpg', 'jpeg', 'png', 'tiff', 'pdf', 'docx']
        
        return result
        
    except Exception as e:
        logger.error(f"Error processing file upload: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
